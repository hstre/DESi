"""Assemble DESi Paper 1 from the existing outline + cited artefacts.

This script does NOT introduce new claims. Every paragraph maps to
an existing artefact already on `release/paper1-representation-gap`.
Citation markers are surfaced as endnotes and audit-listed at the
end of the document.

Constraints:
  - No DESi source change.
  - No new experiments.
  - Only artefacts cited in docs/paper1/outline_representation_gap.md
    (or its appendix A) are pulled.

Output: docs/paper1/DESi_Paper_1_Representation_Gap_v0_1.docx
        docs/paper1/citation_audit_v0_1.md (sibling audit log)
"""
from __future__ import annotations

import json
import pathlib
import re
from collections import defaultdict
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, Inches


# ---------------------------------------------------------------------------
# Artefact registry — every source file the paper is allowed to draw on.
# ---------------------------------------------------------------------------

ARTEFACTS: dict[str, str] = {
    "P0":  "DESi Paper-0 Birth v1-6.docx",
    "BC":  "experiments/birth_criterion_falsification/synthesis.md",
    "ERC": "experiments/external_reality_challenge/synthesis.md",
    "C1E": "experiments/external_reality_challenge/cycle_1/evaluation.md",
    "C1P": "experiments/external_reality_challenge/cycle_1/probe_output_PRE_FIX.json",
    "C2E": "experiments/external_reality_challenge/cycle_2/evaluation.md",
    "C2P": "experiments/external_reality_challenge/cycle_2/probe_output_PRE_FIX.json",
    "PFD": "experiments/external_reality_challenge/post_fix_deltas.md",
    "PPD": "experiments/external_reality_challenge/post_phase_i_ii_fix_deltas.md",
    "EN1": "experiments/external_reality_challenge/en_reconstruction/cycle_1/evaluation.md",
    "EN2": "experiments/external_reality_challenge/en_reconstruction/cycle_2/evaluation.md",
    "EN3": "experiments/external_reality_challenge/en_reconstruction/cycle_3/taxonomy_report.md",
    "EN4R":"experiments/external_reality_challenge/en_reconstruction/cycle_4/downstream_report.md",
    "EN4E":"experiments/external_reality_challenge/en_reconstruction/cycle_4/evaluation.md",
    "EN5R":"experiments/external_reality_challenge/en_reconstruction/cycle_5/time_to_attention_report.md",
    "EN5E":"experiments/external_reality_challenge/en_reconstruction/cycle_5/evaluation.md",
    "RGR": "experiments/external_reality_challenge/representation_gap_report.md",
    "OUT": "docs/paper1/outline_representation_gap.md",
}


@dataclass
class Para:
    text: str
    citations: list[str] = field(default_factory=list)
    style: str = "Body"


@dataclass
class Section:
    number: str
    title: str
    paragraphs: list[Para] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Paper body — each Para is a substantive claim block with attached citations.
# ---------------------------------------------------------------------------

TITLE = "DESi Paper 1 — The Representation Gap"
SUBTITLE = "From Event-Space Diagnostics to Entity-Space Reality"
VERSION = "v0.1 (release/paper1-representation-gap, frozen)"


def build_sections() -> list[Section]:
    sections: list[Section] = []

    # Abstract
    abstract = Section("0", "Abstract")
    abstract.paragraphs.append(Para(
        "Paper 0 established a lower-bound birth criterion for an epistemic "
        "diagnostic system and reported birth(B) = 1 for DESi under the "
        "n=10 adversarial suite.",
        citations=["P0", "BC"],
    ))
    abstract.paragraphs.append(Para(
        "Paper 1 reports what happened when the same DESi was exposed to "
        "its first non-synthetic input: a real DES trajectory dump "
        "(hstre/DES upstream). The system did not crash, but its "
        "diagnostic output was dominated by an interface-failure detector "
        "mislabelling missing data as contradictory data, and by phase-"
        "detector false-positives caused by accidentally-satisfied defaults.",
        citations=["ERC", "C1E", "C2E"],
    ))
    abstract.paragraphs.append(Para(
        "Three targeted interface fixes plus two phase-detector guards "
        "eliminated the false-positive cascade. Two reconstruction rules "
        "then extracted six structurally-verifiable navigation candidates "
        "from native operation history; downstream-effect analysis showed "
        "all six were consequential (5 productive, 1 terminal anchor, 0 "
        "dormant).",
        citations=["PFD", "PPD", "EN1", "EN2", "EN4E"],
    ))
    abstract.paragraphs.append(Para(
        "A time-to-attention comparison between synthetic and native "
        "trajectories proved formally undefined, exposing the underlying "
        "issue: synthetic DESi fixtures encode trajectories in event-space "
        "(discrete EN events with numerical scores), while real DES emits "
        "trajectories in entity-space (claims as named entities with "
        "operations on them). Paper 0's birth criterion was implicitly "
        "event-space. Paper 1 specifies the representation gap, the "
        "interface discipline needed to bridge it, and the open work "
        "needed to satisfy birth(B) under the strong, entity-space-aware "
        "reading.",
        citations=["EN5R", "EN5E", "RGR"],
    ))
    sections.append(abstract)

    # 1. Paper 0 lower-bound birth
    s1 = Section("1", "Paper 0 lower-bound birth")
    s1.paragraphs.append(Para(
        "Paper 0 §12.1 defines B = (D, R, F, P, ΔQ), with birth(B) = 1 iff "
        "all five components are present: failure detection (D), revision "
        "proposal (R), falsifiability (F), preservation of rejected paths "
        "(P), and measurable improvement (ΔQ).",
        citations=["P0"],
    ))
    s1.paragraphs.append(Para(
        "Paper 0 §12.2 reports D = 1, R = 1, F = 1, P = 1, ΔQ > 0 based on "
        "FP: 4 → 0 and FN: 5 → 2 on the n=10 adversarial set, and concludes "
        "birth(B) = 1. §12.3 explicitly frames this as a lower bound, not "
        "an endpoint.",
        citations=["P0"],
    ))
    s1.paragraphs.append(Para(
        "Independent audit on the experimental branch produced per-"
        "component verdicts under both literal and strict readings. Under "
        "Paper 0's literal §12.1 wording all five components hold. Under "
        "stricter readings, D, F, and ΔQ each hold only at the lower bound: "
        "D rests on accepting the same LLM (with different prefix) as "
        "diagnoser and diagnosed; F is satisfied by test-existence rather "
        "than by held-out adversarial conditions; ΔQ is measured on suites "
        "authored by or with knowledge of DESi's failure modes.",
        citations=["BC"],
    ))
    s1.paragraphs.append(Para(
        "The §12.3 lower-bound framing is load-bearing for what follows. "
        "The five components are not independent. D, F, and ΔQ all share "
        "the same hidden assumption that the test-input distribution is "
        "representative of the deployment-input distribution. Paper 1 "
        "exercises that assumption.",
        citations=["P0", "BC"],
    ))
    sections.append(s1)

    # 2. First external DES probe
    s2 = Section("2", "First external DES probe")
    s2.paragraphs.append(Para(
        "Real DES execution state was obtained from `hstre/DES` upstream "
        "main as `des_state.json` and committed verbatim with provenance "
        "notes. The state describes 32 iterations, 9 claims, and 35 "
        "operations. It was authored independent of DESi development; "
        "DESi had never been exposed to it before this probe.",
        citations=["ERC"],
    ))
    s2.paragraphs.append(Para(
        "A translator maps upstream DES schema to DESi trajectory schema "
        "in two declared modes. Conservative mode emits only what upstream "
        "DES contains and leaves DESi-side metrics at their identity "
        "values. Heuristic mode synthesises novel_claims, dup_rate, and "
        "EN events via three declared heuristics, with all synthesis "
        "flagged inline.",
        citations=["ERC", "C1E", "C2E"],
    ))
    s2.paragraphs.append(Para(
        "Under conservative translation, DESi did not crash. The dominant "
        "diagnostic output was step_coherence firing on 34 of 35 steps "
        "with the label \"incoherent\", plus spurious Phase I (medium "
        "confidence) and Phase II (low confidence) triggers, plus an "
        "attractor-detector evaluation that correctly returned no "
        "candidates because the tail was not saturated.",
        citations=["C1E", "C1P"],
    ))
    s2.paragraphs.append(Para(
        "Under heuristic translation, DESi's report shifted: step_coherence "
        "still fired on 17 of 35 steps, eight synthesised EN events were "
        "all classified as `genuine_transformation_unconfirmed`, and Phase "
        "V emitted a high-confidence false-positive span at loops 10–11.",
        citations=["C2E", "C2P"],
    ))
    s2.paragraphs.append(Para(
        "Under the strong reading of components D and ΔQ from Paper 0 "
        "§12.1, the dominant pattern (34/35 misclassified steps plus "
        "phase false-positives) forces birth(B) = 0 on the held-out "
        "distribution. Under §12.1 as literally written (with §12.3's "
        "lower-bound framing), the formal flag still reads 1 — but the "
        "empirical evidence behind the §12.3 disclaimer is now in hand.",
        citations=["ERC", "BC"],
    ))
    sections.append(s2)

    # 3. Schema mismatch
    s3 = Section("3", "Schema mismatch")
    s3.paragraphs.append(Para(
        "Three classes of mismatch were identified between the synthetic "
        "DESi input schema and the native DES output schema, and treated "
        "as the diagnostic interface problem — distinct from any rule-"
        "level diagnostic problem.",
        citations=["ERC"],
    ))
    s3.paragraphs.append(Para(
        "Field-level mismatch. Upstream DES emits no novel_claims, "
        "dup_rate, or eni_* fields. The cycle-6 step_coherence detector "
        "from the self-improvement loop misclassified these absences as "
        "contradictions, producing 34 of 35 false positives in conservative "
        "mode.",
        citations=["ERC", "C1E"],
    ))
    s3.paragraphs.append(Para(
        "Operator-notation mismatch. Real DES emits operation strings of "
        "the form `T6[hypothesis_builder] on C003 -> C008`; DESi's pre-fix "
        "regex matched only the simpler `Tn on Cxxx` form. Six operations "
        "in the sample fell to a silent UNKNOWN substitution, which "
        "downstream cascaded into the heuristic-mode Phase V false-"
        "positive at loops 10–11.",
        citations=["ERC", "C2E"],
    ))
    s3.paragraphs.append(Para(
        "Origin-labelling absence. No input_origin flag existed at the "
        "trajectory level, so reports on translator-derived data were "
        "rendered with the same confidence as reports on hand-authored "
        "data. A reader of the rendered output could not, in principle, "
        "tell whether DESi was diagnosing a fixture or its own "
        "translator's reconstruction.",
        citations=["ERC"],
    ))
    s3.paragraphs.append(Para(
        "DESi did not have a \"trust the input?\" detector at all. The "
        "empirical consequence of this is that every downstream rule "
        "inherited the input ambiguity.",
        citations=["ERC"],
    ))
    sections.append(s3)

    # 4. Interface hardening
    s4 = Section("4", "Interface hardening")
    s4.paragraphs.append(Para(
        "Five targeted fixes were applied to src/desi/. The fixes were "
        "designed to make DESi honest on non-native input without altering "
        "behaviour on the hand-authored suites.",
        citations=["PFD", "PPD"],
    ))
    s4.paragraphs.append(Para(
        "Fix 1 added schema_mismatch detection. TrajectoryStep._normalise "
        "now records which metric fields were absent from the input dict, "
        "and validate_step_metric_coherence skips the dup<0.05/novel=0 "
        "rule when those fields were missing. A new detect_schema_mismatch "
        "report surfaces absent fields per step count and per field.",
        citations=["PFD"],
    ))
    s4.paragraphs.append(Para(
        "Fix 2 introduced parse_des_operation. The new parser handles "
        "both `T3 on C001` and `T6[hypothesis_builder] on C003 -> C008` "
        "forms. On parse failure the parser returns an explicit "
        "OPERATOR_PARSE_FAILURE token rather than a silent UNKNOWN "
        "substitution.",
        citations=["PFD"],
    ))
    s4.paragraphs.append(Para(
        "Fix 3 added an input_origin field on Trajectory with canonical "
        "values hand_authored_fixture, translator_heuristic, live_DES, "
        "and translated_DES_conservative. render_report prepends a "
        "translator-derived disclaimer when input_origin is non-"
        "authoritative.",
        citations=["PFD"],
    ))
    s4.paragraphs.append(Para(
        "Fix 4 made detect_phase_i silent when loop 0's metrics are "
        "missing. Pre-fix the detector emitted medium-confidence Phase I "
        "because the defaulted-to-0 dup_rate accidentally satisfied the "
        "<0.30 condition. Post-fix the detector treats the missing-metric "
        "condition as INSUFFICIENT_METRICS and returns None.",
        citations=["PPD"],
    ))
    s4.paragraphs.append(Para(
        "Fix 5 made detect_phase_ii silent when novel_claims is missing "
        "across the persistence pair. Pre-fix the persistence rule was "
        "trivially satisfied by defaulted-zero novel claims on every "
        "consecutive step pair. Post-fix the pair is skipped when either "
        "side has novel_claims in missing_metrics.",
        citations=["PPD"],
    ))
    s4.paragraphs.append(Para(
        "The five-fix package raises pytest from 35 to 49. On the n=10 "
        "adversarial suite and the n=20 generalization suite, every "
        "detector fire count is unchanged. On the external DES "
        "conservative-mode probe, step_coherence drops from 34/35 to "
        "0/35; both spurious Phase I and Phase II disappear; "
        "schema_mismatch reports True with 35/35 missing fields; phases "
        "= [].",
        citations=["PFD", "PPD"],
    ))
    s4.paragraphs.append(Para(
        "Interface hardening is necessary but not sufficient. The fixes "
        "stop DESi from mislabelling non-native input, but they do not "
        "give DESi the ability to diagnose the underlying DES trajectory. "
        "The remainder of the paper asks whether a useful diagnostic "
        "ceiling exists.",
        citations=["PFD", "PPD"],
    ))
    sections.append(s4)

    # 5. Reconstruction of structural navigation candidates
    s5 = Section("5", "Reconstruction of structural navigation candidates")
    s5.paragraphs.append(Para(
        "Two reconstruction rules emit candidates from native DES "
        "operation history alone. Neither rule fabricates an eni_novelty "
        "score.",
        citations=["EN1", "EN2"],
    ))
    s5.paragraphs.append(Para(
        "Rule 1 (EN candidate): the candidate fires iff "
        "operator_sub_role equals \"hypothesis_builder\" and "
        "operator_target is not None. The rule emits LOCATIONS — loop "
        "indices and created-claim identities — not numerical scores.",
        citations=["EN1"],
    ))
    s5.paragraphs.append(Para(
        "Rule 2 (critique-navigation candidate): the candidate fires iff "
        "operator_sub_role equals \"falsifier\" and operator_target is "
        "not None. The two rules' candidate sets are kept disjoint by "
        "design; per-step the operator carries at most one sub-role, so "
        "the loop indices cannot overlap.",
        citations=["EN2"],
    ))
    s5.paragraphs.append(Para(
        "On real DES the two rules fire three times each (six candidates "
        "total). EN candidates are at loops 3, 7, 10. Critique-navigation "
        "candidates are at loops 4, 8, 11. The two sets are disjoint.",
        citations=["EN1", "EN2", "EN3"],
    ))
    s5.paragraphs.append(Para(
        "A 5-way taxonomy partitions all 35 upstream operations: 3 EN "
        "candidates (8.6%), 3 critique-navigation candidates (8.6%), 29 "
        "plain operator transitions (82.9%), 0 unsupported extensions, "
        "and 0 unparsed operations. The taxonomy uses parse_des_operation "
        "and the two rules; it does not introduce a new rule.",
        citations=["EN3"],
    ))
    s5.paragraphs.append(Para(
        "Three measurements follow from the taxonomy. The unparsed rate "
        "is 0.0%; the unsupported rate is 0.0%; and the target-creating "
        "completeness is 6/6 = 100.0%. That last figure is the strongest "
        "claim the reconstruction layer supports: every operation in "
        "upstream DES that produces a new target claim falls into a "
        "reconstruction category, with no slippage into "
        "unsupported_extension or plain_operator_transition.",
        citations=["EN3"],
    ))
    s5.paragraphs.append(Para(
        "Reconstruction does not recover the ENI signal. It recovers "
        "locations where DES extended its claim graph via a typed "
        "operation. The two categories are kept separate by design; "
        "merging them would have erased a finding documented in the next "
        "section.",
        citations=["EN1", "EN2"],
    ))
    sections.append(s5)

    # 6. Downstream validation
    s6 = Section("6", "Downstream validation")
    s6.paragraphs.append(Para(
        "For each of the six created claims (C004 through C009) the "
        "downstream-effect analyser computes four signals from native DES "
        "fields only: first_focus_after_creation, "
        "touch_count_after_creation, survives_in_final_claims, and "
        "is_terminal_focus.",
        citations=["EN4R", "EN4E"],
    ))
    s6.paragraphs.append(Para(
        "Candidates are then classified into three mutually exclusive "
        "categories: terminal_anchor when the claim equals the final "
        "focus_claim_id; productive when not-terminal AND the touch count "
        "is at least three AND the claim survives in the final claims "
        "dict; dormant otherwise.",
        citations=["EN4R", "EN4E"],
    ))
    s6.paragraphs.append(Para(
        "The result on real DES is 5 productive plus 1 terminal_anchor "
        "plus 0 dormant. Every reconstructed candidate became something "
        "DES continued to work on for the rest of the trajectory.",
        citations=["EN4R", "EN4E"],
    ))
    s6.paragraphs.append(Para(
        "Four observations follow. First, no candidate is dormant: 0/6. "
        "Second, mean creation-to-first-focus latency on real DES is "
        "approximately 15 loops, an order of magnitude longer than the "
        "single-digit timescales hand-authored fixtures assume. Third, "
        "the terminal anchor (C009) was created via a `falsifier` "
        "operation, not via `hypothesis_builder`; the trajectory ended on "
        "a critique-derived claim. Fourth, the six candidates form three "
        "hypothesis+falsifier pairs operating on the same parent claim "
        "in adjacent loops, suggesting a paired-move pattern that could "
        "be reconstructed as a higher-order rule in future work.",
        citations=["EN4E"],
    ))
    s6.paragraphs.append(Para(
        "Three caveats apply. The analysis is run on n=1 trajectory; "
        "generalisation requires replication on additional dumps. "
        "Productivity is correlation with downstream activity, not "
        "causation. The threshold PRODUCTIVE_MIN_TOUCHES = 3 is a "
        "parameter of the analysis, documented explicitly, not embedded "
        "as a DESi-side constant.",
        citations=["EN4E"],
    ))
    sections.append(s6)

    # 7. Time-to-attention undefined comparison
    s7 = Section("7", "Time-to-attention: an undefined comparison")
    s7.paragraphs.append(Para(
        "For every reconstructed candidate the latency metric "
        "latency = first_focus_loop − creation_loop was computed across "
        "three sources: the hand-authored n=10 adversarial suite, the "
        "hand-authored n=20 generalization suite, and the native DES "
        "trajectory.",
        citations=["EN5R", "EN5E"],
    ))
    s7.paragraphs.append(Para(
        "On the hand-authored suites the reconstruction rules emit zero "
        "candidates because no fixture in either suite uses the "
        "[hypothesis_builder] or [falsifier] sub-role notation. The "
        "latency metric therefore has no synthetic-side sample. On real "
        "DES the rules emit six candidates with mean latency 15.17, "
        "median 15.0, and variance 12.97.",
        citations=["EN5R"],
    ))
    s7.paragraphs.append(Para(
        "The pre-defined flag LATENCY_DISTRIBUTION_MISMATCH = TRUE iff "
        "real_DES_mean > 2 × synthetic_mean evaluates to UNDEFINED, not "
        "TRUE or FALSE. A ratio with no denominator is not a number. "
        "Treating synthetic_mean as zero would force the flag TRUE for "
        "any non-zero real_mean; treating it as undefined-equal-to-zero "
        "is a category error.",
        citations=["EN5E"],
    ))
    s7.paragraphs.append(Para(
        "Two compounding reasons explain the undefinedness. First, the "
        "reconstruction rules require sub-role annotations that hand-"
        "authored fixtures do not carry; the candidate set is empty. "
        "Second, even if synthetic fixtures had sub-role-tagged "
        "operations, the metric needs the created claim to become a "
        "focus_claim_id later; hand-authored fixtures hold focus_claim_id "
        "functionally constant.",
        citations=["EN5E"],
    ))
    s7.paragraphs.append(Para(
        "The formal undefinedness is more informative than a TRUE/FALSE "
        "answer would have been. It shows that the two corpora do not "
        "share a metric framework. The qualitative finding — real-DES "
        "latency far exceeds the implicit single-digit timescale of the "
        "hand-authored fixtures — remains; only the formal flag is "
        "undefined.",
        citations=["EN5E"],
    ))
    sections.append(s7)

    # 8. Event-space vs entity-space
    s8 = Section("8", "Event-space vs entity-space")
    s8.paragraphs.append(Para(
        "The undefinedness in §7 reflects a difference in representation "
        "space, not a difference in dataset size or trajectory shape.",
        citations=["RGR"],
    ))
    s8.paragraphs.append(Para(
        "Synthetic DESi fixtures operate in event-space. The primary unit "
        "of analysis is the EN event: a discrete moment with a numerical "
        "novelty score and pre/post duplication snapshots. Per-claim "
        "identity over time is not tracked. Across all thirty hand-"
        "authored fixtures focus_claim_id is functionally constant. "
        "Roughly 8–10 distinct numerical signals are carried per EN event.",
        citations=["RGR"],
    ))
    s8.paragraphs.append(Para(
        "Native DES operates in entity-space. The primary unit of "
        "analysis is the claim — a named entity that the trajectory "
        "creates, modifies, and ultimately converges onto. There are no "
        "native ENI scores. There is no en_events list. There is per-"
        "claim identity over time: C001 through C009 are distinct nodes "
        "the trajectory operates on. Roughly seven distinct signals are "
        "carried per graph-extending operation.",
        citations=["RGR"],
    ))
    s8.paragraphs.append(Para(
        "Reading Paper 0 with this gap in view exposes four implicit "
        "assumptions. (A) Trajectories arrive in event-space — i.e. the "
        "upstream component measures novelty, duplication, and entropy-"
        "navigation events and ships them as fields to DESi. (B) "
        "focus_claim_id is a label, not a tracked entity. (C) Recovery is "
        "measured at event boundaries via dup_delta plus novel_claims_next. "
        "(D) The birth criterion B = (D, R, F, P, ΔQ) is measurable in "
        "event-space.",
        citations=["RGR", "P0"],
    ))
    s8.paragraphs.append(Para(
        "The empirical facts contradict all four. DES does not measure "
        "novelty generation or duplication dynamics; real DES tracks "
        "focus_claim_id as an evolving entity; there are no discrete "
        "pre/post snapshots in real DES output; and when DESi's rules "
        "are exercised on real DES, they go silent. Paper 0's ΔQ "
        "measurement (FP 4 → 0; FN 5 → 2) was an event-space-only result.",
        citations=["RGR", "PFD", "PPD"],
    ))
    sections.append(s8)

    # 9. Requirements for DESi entity-space
    s9 = Section("9", "Requirements for DESi entity-space")
    s9.paragraphs.append(Para(
        "If a DESi-on-real-DES system is to satisfy a strengthened birth "
        "criterion, five requirements must be addressed explicitly.",
        citations=["RGR"],
    ))
    s9.paragraphs.append(Para(
        "R1 — Define ENI in entity-space, or admit it does not transfer. "
        "Three defensible options: lift the measurement to DES (requires "
        "DES-side change); define an entity-space proxy and falsify it; "
        "or formally separate DESi-event-space from DESi-entity-space as "
        "two systems sharing a codebase.",
        citations=["RGR"],
    ))
    s9.paragraphs.append(Para(
        "R2 — Define recovery in entity-space. The cycle-4 "
        "productive/dormant/terminal_anchor classification is one "
        "candidate operational predicate. Paper 1 should commit to a "
        "formulation and provide a falsification probe.",
        citations=["RGR", "EN4E"],
    ))
    s9.paragraphs.append(Para(
        "R3 — Re-derive the five birth(B) components on entity-space. For "
        "each of D, R, F, P, ΔQ, state the entity-space meaning, the "
        "measurement protocol, and what evidence on real DES would force "
        "the component to 0. The current Paper 0 §12.2 evaluation cannot "
        "be ported wholesale because the FP/FN counters themselves are "
        "event-space-defined.",
        citations=["RGR", "P0", "BC"],
    ))
    s9.paragraphs.append(Para(
        "R4 — Acknowledge that \"DESi diagnoses DES\" is two systems. "
        "DESi-event-space is mature (~58 tests, ingests hand-authored "
        "EN-bearing trajectories, rich rule set). DESi-entity-space is "
        "immature (a parser, two reconstruction rules, a schema-mismatch "
        "detector, downstream-effect analyser; no phase, attractor, or "
        "failure-mode diagnostics). They are not two implementations of "
        "one system; they are two systems with overlapping codebases.",
        citations=["RGR"],
    ))
    s9.paragraphs.append(Para(
        "R5 — Stop reporting in-distribution metrics as evidence of "
        "generalisation. Paper 0 §11.10 acknowledges that broader "
        "generalisation across natural domains, long trajectories, and "
        "semantic complexity remains unestablished. The metric "
        "improvements of the self-improvement and generalisation loops "
        "are improvements on a representation; they are not improvements "
        "of the system on its intended target.",
        citations=["RGR", "P0"],
    ))
    sections.append(s9)

    # 10. Open work
    s10 = Section("10", "Open work")
    s10.paragraphs.append(Para(
        "Seven open items are inherited from the experimental branch's "
        "outline and from the cycle-4 caveats. None are proposed as new "
        "experiments inside this paper.",
        citations=["OUT"],
    ))
    s10.paragraphs.append(Para(
        "O1 — Replication on additional real DES dumps. The "
        "5 productive + 1 terminal_anchor + 0 dormant result rests on "
        "n = 1 trajectory.",
        citations=["OUT", "EN4E"],
    ))
    s10.paragraphs.append(Para(
        "O2 — Entity-space ENI formula. Requirement R1 from §9 cannot be "
        "resolved from existing artefacts; the existing reconstruction "
        "cycles emit locations, not scores.",
        citations=["OUT", "RGR"],
    ))
    s10.paragraphs.append(Para(
        "O3 — Entity-space birth(B) re-derivation. The audit covers the "
        "event-space reading of D/R/F/P/ΔQ; the entity-space reading is "
        "unaddressed.",
        citations=["OUT", "BC", "RGR"],
    ))
    s10.paragraphs.append(Para(
        "O4 — Paired-builder-falsifier as a higher-order signal. The "
        "pattern (cycle 4 observation O4) was noted but not implemented "
        "as a detector.",
        citations=["OUT", "EN4E"],
    ))
    s10.paragraphs.append(Para(
        "O5 — Cross-validation against entity-space-aware hand-authored "
        "fixtures. The current hand-authored suites do not exercise "
        "operator_sub_role or operator_target.",
        citations=["OUT"],
    ))
    s10.paragraphs.append(Para(
        "O6 — Falsifier-terminal-anchor case study. The cycle-4 finding "
        "(O3) that C009 came from a `falsifier` is a small, citable "
        "result that defends the cycle-2 disjointness directive "
        "empirically.",
        citations=["OUT", "EN4E", "EN2"],
    ))
    s10.paragraphs.append(Para(
        "O7 — Taxonomy run on the upstream `des_prototype/des_state.json`. "
        "The prototype state file is committed but no cycle has classified "
        "it; a single additional run would either confirm cycle 3's "
        "100% target-creating completeness or expose a counterexample.",
        citations=["OUT", "EN3"],
    ))
    sections.append(s10)

    return sections


# ---------------------------------------------------------------------------
# Citation audit — flag uncited claims and duplicated claims.
# ---------------------------------------------------------------------------


def audit_citations(sections: list[Section]) -> dict:
    audit = {
        "total_paragraphs": 0,
        "uncited_paragraphs": [],     # (section, idx, text-preview)
        "duplicated_text": [],        # ((section_a, idx_a), (section_b, idx_b))
        "artefact_use_counts": defaultdict(int),
        "unused_artefacts": [],
    }
    # Track first-occurrence to detect duplicates.
    text_to_loc: dict[str, tuple[str, int]] = {}
    for sec in sections:
        for idx, p in enumerate(sec.paragraphs):
            audit["total_paragraphs"] += 1
            if not p.citations:
                audit["uncited_paragraphs"].append(
                    (sec.number, idx, p.text[:80])
                )
            for cit in p.citations:
                audit["artefact_use_counts"][cit] += 1
            # Normalise text for duplicate detection: lowercase, collapse ws.
            key = re.sub(r"\s+", " ", p.text.lower()).strip()
            if key in text_to_loc:
                audit["duplicated_text"].append(
                    (text_to_loc[key], (sec.number, idx))
                )
            else:
                text_to_loc[key] = (sec.number, idx)
    for code in ARTEFACTS:
        if audit["artefact_use_counts"][code] == 0:
            audit["unused_artefacts"].append(code)
    return audit


def write_audit_markdown(audit: dict, out_path: pathlib.Path) -> None:
    lines = ["# Paper 1 v0.1 — citation audit",
             "",
             f"Total paragraphs: **{audit['total_paragraphs']}**",
             ""]
    lines.append("## Per-artefact use counts")
    lines.append("")
    lines.append("| Code | Artefact | Uses |")
    lines.append("|---|---|---:|")
    for code, path in ARTEFACTS.items():
        n = audit["artefact_use_counts"].get(code, 0)
        lines.append(f"| `{code}` | `{path}` | {n} |")
    lines.append("")
    lines.append("## Uncited paragraphs")
    lines.append("")
    if not audit["uncited_paragraphs"]:
        lines.append("- None. Every paragraph carries at least one citation.")
    else:
        for sec, idx, preview in audit["uncited_paragraphs"]:
            lines.append(f"- §{sec} ¶{idx+1}: \"{preview}…\"")
    lines.append("")
    lines.append("## Duplicated text spans")
    lines.append("")
    if not audit["duplicated_text"]:
        lines.append(
            "- None. No two paragraphs share normalised text under "
            "case-insensitive whitespace-collapsed comparison."
        )
    else:
        for (a_sec, a_i), (b_sec, b_i) in audit["duplicated_text"]:
            lines.append(f"- §{a_sec} ¶{a_i+1} ≡ §{b_sec} ¶{b_i+1}")
    lines.append("")
    lines.append("## Unused artefacts")
    lines.append("")
    if not audit["unused_artefacts"]:
        lines.append("- None. Every registered artefact is cited at least once.")
    else:
        for code in audit["unused_artefacts"]:
            lines.append(f"- `{code}` (`{ARTEFACTS[code]}`)")
    lines.append("")
    out_path.write_text("\n".join(lines))


# ---------------------------------------------------------------------------
# Docx writer.
# ---------------------------------------------------------------------------


def write_docx(sections: list[Section], audit: dict, out_path: pathlib.Path) -> None:
    doc = Document()
    # A light style: 11pt body, 14pt section headings, 18pt title.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(TITLE, level=0)
    p = doc.add_paragraph()
    r = p.add_run(SUBTITLE)
    r.italic = True
    p2 = doc.add_paragraph(f"Version: {VERSION}")
    p2.runs[0].italic = True

    doc.add_paragraph(
        "This document is assembled from existing artefacts on the "
        "release/paper1-representation-gap branch. Every paragraph carries "
        "citation markers (in square brackets) keyed to Appendix A. No new "
        "experiments were run to produce this paper. Citation audit, "
        "uncited-claim flags, and duplicate-claim flags appear in the "
        "sibling file citation_audit_v0_1.md."
    )

    for sec in sections:
        if sec.number == "0":
            doc.add_heading("Abstract", level=1)
        else:
            doc.add_heading(f"{sec.number}. {sec.title}", level=1)
        for p in sec.paragraphs:
            para = doc.add_paragraph(p.text)
            if p.citations:
                # Append a small bracketed cite tag, separated by a tab.
                run = para.add_run(" [" + ", ".join(p.citations) + "]")
                run.italic = True
                run.font.size = Pt(9)

    # Appendix A — artefact key
    doc.add_heading("Appendix A — Artefact key", level=1)
    doc.add_paragraph(
        "All paths are relative to the repository root on "
        "release/paper1-representation-gap."
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Code"
    hdr[1].text = "Artefact path"
    for code, path in ARTEFACTS.items():
        row = tbl.add_row().cells
        row[0].text = code
        row[1].text = path

    # Appendix B — audit summary
    doc.add_heading("Appendix B — Citation audit summary", level=1)
    doc.add_paragraph(
        f"Total paragraphs: {audit['total_paragraphs']}. "
        f"Uncited paragraphs: {len(audit['uncited_paragraphs'])}. "
        f"Duplicated paragraph spans: {len(audit['duplicated_text'])}. "
        f"Unused artefacts: {len(audit['unused_artefacts'])}. "
        "Full detail in citation_audit_v0_1.md."
    )

    # Appendix C — constraints
    doc.add_heading("Appendix C — Constraints honoured", level=1)
    doc.add_paragraph(
        "(1) No DESi source change in src/desi/ for this release; pytest "
        "58 of 58 unchanged at HEAD. (2) No new experiments were run to "
        "produce this paper. (3) Only artefacts listed in Appendix A are "
        "cited. (4) The branch release/paper1-representation-gap is "
        "frozen for v0.1; subsequent revisions create a new versioned "
        "file rather than overwrite this one."
    )

    doc.save(out_path)


# ---------------------------------------------------------------------------
# Main.
# ---------------------------------------------------------------------------


def main() -> int:
    # Verify all cited artefacts exist on disk.
    repo_root = pathlib.Path(__file__).resolve().parents[2]
    missing = [
        code for code, path in ARTEFACTS.items()
        if not (repo_root / path).is_file()
    ]
    if missing:
        print(f"MISSING ARTEFACTS: {missing}", flush=True)
        return 2
    sections = build_sections()
    audit = audit_citations(sections)
    out_dir = repo_root / "docs" / "paper1"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "DESi_Paper_1_Representation_Gap_v0_1.docx"
    audit_path = out_dir / "citation_audit_v0_1.md"
    write_docx(sections, audit, docx_path)
    write_audit_markdown(audit, audit_path)
    print(f"wrote {docx_path}")
    print(f"wrote {audit_path}")
    print(f"paragraphs={audit['total_paragraphs']}  "
          f"uncited={len(audit['uncited_paragraphs'])}  "
          f"duplicated={len(audit['duplicated_text'])}  "
          f"unused_artefacts={len(audit['unused_artefacts'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
